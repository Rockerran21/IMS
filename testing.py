from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_custom_preformatted_style(doc):
    styles = doc.styles
    preformatted_style = styles.add_style('PreformattedText', 1)  # 1 corresponds to WD_STYLE_TYPE.PARAGRAPH
    font = preformatted_style.font
    font.name = 'Courier New'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = 'PreformattedText'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10.5)

def create_architecture_doc(file_path):
    # Create a Document
    doc = Document()

    # Title
    title = doc.add_heading('Architecture Overview', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('This document provides an architecture diagram and detailed explanation of the key components involved in deploying the Llama 3.1 model in an Azure Kubernetes Service (AKS) cluster.')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Architecture Diagram
    doc.add_heading('Architecture Diagram', level=2)
    # Insert Diagram Image (placeholder)
    doc.add_picture('/Users/ranjanmarasini/IdeaProjects/IMS/RA for ChatBot Application.png', width=Pt(500))  # Adjust the path and width accordingly

    # Components
    doc.add_heading('Components', level=2)
    components = [
        ('Azure Kubernetes Service (AKS)', 'The central component that orchestrates the deployment, scaling, and management of Docker containers.'),
        ('Docker', 'Used for containerizing the Ollama model.'),
        ('Ollama Model Container', 'The Docker container that runs the Ollama model.'),
        ('Azure Container Registry (ACR)', 'The repository for storing and managing Docker container images.'),
        ('Azure Virtual Network (VNet)', 'Provides network connectivity between AKS and other Azure resources.'),
        ('Load Balancer', 'Distributes incoming network traffic across multiple instances of the Ollama model container.'),
        ('Azure Key Vault', 'Securely stores and manages sensitive information such as keys, secrets, and certificates.'),
        ('Ingress Controller', 'Manages external access to the services in the AKS cluster, usually HTTP and HTTPS routes.'),
        ('End User', 'Interacts with the Chatbot.'),
        ('Chatbot', 'Requests model predictions from AKS.')
    ]

    for comp, desc in components:
        p = doc.add_paragraph()
        p.add_run(f'{comp}: ').bold = True
        p.add_run(desc)

    # Steps to Deploy the Llama 3.1 Model on AKS
    doc.add_heading('Steps to Deploy the Llama 3.1 Model on AKS', level=2)
    steps = [
        ('Setup Azure Kubernetes Service (AKS)', 'Create an AKS cluster in the Azure portal.', '''
        az aks create --resource-group myResourceGroup --name myAKSCluster --node-count 1 --enable-addons monitoring --generate-ssh-keys
        '''),
        ('Build and Push Docker Image', 'Build a Docker image for the Ollama model and push it to Azure Container Registry (ACR).', '''
        # Dockerfile
        FROM ubuntu:latest
        RUN apt-get update && apt-get install -y python3-pip
        COPY . /app
        WORKDIR /app
        RUN pip3 install -r requirements.txt
        CMD ["python3", "app.py"]

        # Build and Push
        docker build -t ollama-model .
        az acr login --name myACR
        docker tag ollama-model myacr.azurecr.io/ollama-model:v1
        docker push myacr.azurecr.io/ollama-model:v1
        '''),
        ('Deploy the Model to AKS', 'Create a Kubernetes deployment and service for the Ollama model.', '''
        # deployment.yaml
        apiVersion: apps/v1
        kind: Deployment
        metadata:
          name: ollama-model
        spec:
          replicas: 3
          selector:
            matchLabels:
              app: ollama-model
          template:
            metadata:
              labels:
                app: ollama-model
            spec:
              containers:
              - name: ollama-model
                image: myacr.azurecr.io/ollama-model:v1
                ports:
                - containerPort: 80

        # service.yaml
        apiVersion: v1
        kind: Service
        metadata:
          name: ollama-model
        spec:
          selector:
            app: ollama-model
          ports:
            - protocol: TCP
              port: 80
              targetPort: 80
          type: LoadBalancer

        # Apply the configurations
        kubectl apply -f deployment.yaml
        kubectl apply -f service.yaml
        '''),
        ('Configure Ingress Controller', 'Setup an Ingress Controller to manage external access to the services in the AKS cluster.', '''
        # ingress-controller.yaml
        apiVersion: networking.k8s.io/v1
        kind: Ingress
        metadata:
          name: ollama-ingress
        spec:
          rules:
          - host: ollama.example.com
            http:
              paths:
              - path: /
                pathType: Prefix
                backend:
                  service:
                    name: ollama-model
                    port:
                      number: 80

        # Apply the ingress configuration
        kubectl apply -f ingress-controller.yaml
        '''),
        ('Setup Azure Key Vault', 'Store and manage sensitive information such as keys, secrets, and certificates.', '''
        # Create Key Vault
        az keyvault create --name myKeyVault --resource-group myResourceGroup --location eastus

        # Store secrets
        az keyvault secret set --vault-name myKeyVault --name "MySecret" --value "mySecretValue"
        '''),
        ('Connect AKS to Key Vault', 'Configure AKS to access secrets from Azure Key Vault.', '''
        # Assign managed identity to AKS
        az aks update -g myResourceGroup -n myAKSCluster --enable-managed-identity

        # Assign Key Vault access policy
        az keyvault set-policy -n myKeyVault --secret-permissions get --spn <your-aks-managed-identity-client-id>
        '''),
        ('Deploy Chatbot', 'Deploy the chatbot application that interacts with end-users and requests model predictions from AKS.', '''
        # deployment.yaml for chatbot
        apiVersion: apps/v1
        kind: Deployment
        metadata:
          name: chatbot
        spec:
          replicas: 2
          selector:
            matchLabels:
              app: chatbot
          template:
            metadata:
              labels:
                app: chatbot
            spec:
              containers:
              - name: chatbot
                image: myacr.azurecr.io/chatbot:v1
                ports:
                - containerPort: 8080

        # service.yaml for chatbot
        apiVersion: v1
        kind: Service
        metadata:
          name: chatbot
        spec:
          selector:
            app: chatbot
          ports:
            - protocol: TCP
              port: 8080
              targetPort: 8080
          type: LoadBalancer

        # Apply the configurations
        kubectl apply -f chatbot-deployment.yaml
        kubectl apply -f chatbot-service.yaml
        ''')
    ]

    for title, description, code in steps:
        doc.add_heading(title, level=3)
        doc.add_paragraph(description)
        doc.add_paragraph(code, style='Preformatted Text')

    # Save Document
    doc.save(file_path)

# Create the document
create_architecture_doc("Architecture_Overview.docx")
